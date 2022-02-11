import getpass
from io import StringIO
from random import randint
from datetime import date
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.utils.encoding import uri_to_iri
from django.db.models import Max
from django.contrib.auth import authenticate, login
from django.contrib.auth.models import Group, User
from django.shortcuts import render, HttpResponse
#from transliterate import slugify
from docx import Document
from docx.shared import Pt
from .models import *
from .middleware import check_network, GetDataFromAD
import pylightxl as xl
from openpyxl import Workbook
from itertools import islice
from django.templatetags.static import static
import os
from django.conf import settings


def word_1kk(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)

        mos = MunObr.objects.all()
        count1kk = 1
        file = open(settings.STATIC_ROOT + '\\projects\\Proj1KK.docx', 'rb')
        doc_1kk = Document(file)
        style = doc_1kk.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        for mo in mos:
            one_kk = certified.objects.filter(period_id=period).filter(MO_id=mo.id).filter(result='Установить').\
                filter(Category_id=AttCategories.objects.get(type_cat='первая').id)
            if one_kk.count() > 0:
                name = mo.name_MO
                name = name[:1].upper() + name[1:]
                doc_1kk.add_paragraph(name+":")
                for cert in one_kk:
                    oo = cert.Organization
                    oo = oo[:1].lower() + oo[1:]
                    doc_1kk.add_paragraph(str(count1kk)+'. '+cert.FIO+', '+str(cert.Position).lower()+', '+oo+';')
                    count1kk += 1
        response = HttpResponse(doc_1kk)
        response['Content-Disposition'] = 'attachment; filename=1KK.docx'
        doc_1kk.save(response)
        return response
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def word_vkk(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)

        mos = MunObr.objects.all()
        countvkk = 1
        file = open(settings.STATIC_ROOT + '\\projects\\ProjVKK.docx', 'rb')
        doc_vkk = Document(file)
        style = doc_vkk.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        for mo in mos:
            vkk = certified.objects.filter(period_id=period).filter(MO_id=mo.id).filter(result='Установить').\
                filter(Category_id=AttCategories.objects.get(type_cat='высшая').id)
            if vkk.count() > 0:
                name = mo.name_MO
                name = name[:1].upper() + name[1:]
                doc_vkk.add_paragraph(name.capitalize()+":")
                for cert in vkk:
                    oo = cert.Organization
                    oo = oo[:1].lower() + oo[1:]
                    doc_vkk.add_paragraph(str(countvkk)+'. '+cert.FIO+', '+str(cert.Position).lower()+', '+oo+';')
                    countvkk += 1
        response = HttpResponse(doc_vkk)
        response['Content-Disposition'] = 'attachment; filename=VKK.docx'
        doc_vkk.save(response)
        return response
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def get_user_and_group(request):
    try:
        from_ad = GetDataFromAD(request)
        username = from_ad[0][0]
        group = from_ad[1][0]
        email = from_ad[2][0]
        if check_network(request) == False:
            username = request.user.username
            group = 'Аттестация'
            email = request.user.email
    except BaseException:
        username = request.user.username
        list_group = request.user.groups.values_list('name', flat=True)
        group = list_group[0]
        email = request.user.email
    return [username, group, email]


def check_access(request, group):
    if group == 'Уполномоченный' or group == 'Специалист' or check_network(request) == False:
        return True
    else:
        return False
        

def change_mail(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        mail = data[2]
        if request.method == 'POST':
            us = request.user
            us.email = request.POST['email']
            us.save()
            context = {
                'curr_user': user,
                'curr_group': group,
                'curr_mail': mail
            }
            return render(request, 'start.html', context)
        else:
            context = {
                'curr_user': user,
                'curr_group': group,
                'curr_mail': mail
            }
            return render(request, 'profile.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def user_profile(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        mail = data[2]
        context = {
            'curr_user': user,
            'curr_group': group,
            'curr_mail': mail
        }
        return render(request, 'profile.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def update_all(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        ExpCards.objects.all().delete()
        criteria_export.objects.all().delete()
        interm_recs = intermediate.objects.all()
        for el in interm_recs:
            inter_id = el.id
            spec = specialists.objects.get(id=el.spec_id)
            att = certified.objects.get(id=el.cert_id)
            operator_choice = el.operator_choice
            new_card = ExpCards()
            new_card.period_id = spec.period_id
            new_card.date_add = el.date_add
            if operator_choice == 'Специалист':
                new_card.type_emp = operator_choice
            else:
                new_card.type_emp = 'Оператор'
            id_mo_spec = spec.MO_id
            type_mo = MunObr.objects.get(id=id_mo_spec).type_MO
            new_card.type_MO = type_mo
            new_card.MO_expert_id = spec.MO_id
            new_card.FIO_expert = spec.FIO
            new_card.Name_Org_expert = spec.Organization
            new_card.Position_expert_id = spec.Position_id
            new_card.Level_expert_id = el.level_id
            new_card.MO_att_id = att.MO_id
            new_card.FIO_att = att.FIO
            new_card.Name_Org_att = att.Organization
            new_card.Position_att_id = att.Position_id
            new_card.Category_id = att.Category_id
            new_card.Result = el.result
            new_card.inter_id = inter_id
            new_card.save()
        crit_recs = cards_criteria.objects.all()
        for el in crit_recs:
            spec = cards_criteria.objects.get(id=el.id).spec_id
            cert = cards_criteria.objects.get(id=el.id).cert_id
            new_crit = criteria_export()
            new_crit.period_id = specialists.objects.get(id=spec).period_id
            new_crit.id_cards_criteria_id = el.id
            new_crit.MO_spec_id = specialists.objects.get(id=spec).MO_id
            new_crit.FIO_spec = specialists.objects.get(id=spec).FIO
            new_crit.Position_spec_id = specialists.objects.get(id=spec).Position_id
            new_crit.MO_att_id = certified.objects.get(id=cert).MO_id
            new_crit.FIO_att = certified.objects.get(id=cert).FIO
            new_crit.Position_att_id = certified.objects.get(id=cert).Position_id
            new_crit.Category_id = certified.objects.get(id=cert).Category_id
            new_crit.Result = intermediate.objects.get(id=el.inter_id).result
            new_crit.criteria_id = el.criteria_id
            new_crit.info = el.info
            new_crit.save() 
        return SumTab(request)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def SumTab(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            Summary_table.objects.all().delete()
            per = request.POST.get('period')
            #Поиск недостающих экспертов, которых нет в итоговой таблице
            obj = specialists.objects.filter(period_id=per).select_related('period')
            list_spec = []
            for el in obj:
                list_spec.append(el.id)
            specs = intermediate.objects.filter(spec_id__in=list_spec)\
                .values('spec_id', 'level_id', 'operator_choice').distinct()
            string = ''
            for el in specs:
                spec = specialists.objects.get(id=el['spec_id'])
                check_new = False
                try:
                    res = Summary_table.objects.filter(FIO_expert=spec.FIO).filter(MO_expert_id=spec.MO_id) \
                        .filter(Position_expert_id=spec.Position_id).latest('id')
                except BaseException:
                    check_new = True
                if check_new:
                    if el['operator_choice'] == 'Специалист':
                        count_common = 0
                        count_general = 0
                        count_operator = 0
                        for el_int in intermediate.objects.filter(spec_id__in=list_spec):
                            if el_int.spec_id == spec.id and el_int.level_id == el['level_id']:
                                for el_subint in intermediate.objects.filter(spec_id__in=list_spec):
                                    if el_int.cert_id == el_subint.cert_id and el_int.spec_id != el_subint.spec_id \
                                            and el_int.result == el_subint.result:
                                        if el_subint.level_id == ExpLevel.objects.get(name_level='Оператор').id:
                                            count_operator += 1
                                        elif el_subint.level_id == ExpLevel.objects.get(name_level='Первый этап').id:
                                            check_right = True
                                            for el_third in intermediate.objects.filter(spec_id__in=list_spec):
                                                if el_int.cert_id == el_third.cert_id \
                                                        and el_int.level_id != el_third.level_id and \
                                                        el_int.result != el_third.result:
                                                    check_right = False
                                            if check_right:
                                                count_common += 1
                                        else:
                                            count_general += 1
                        cnt = intermediate.objects.filter(spec_id=spec.id).filter(level_id=el['level_id']).count()
                        new_rec = Summary_table(
                            period_id=spec.period_id,
                            FIO_expert=spec.FIO,
                            Level_expert_id=el['level_id'],
                            MO_expert_id=spec.MO_id,
                            Position_expert_id=spec.Position_id,
                            count=cnt,
                            coincidence_common=count_common,
                            coincidence_general=count_general,
                            coincidence_operator=count_operator,
                            type_MO=MunObr.objects.get(id=spec.MO_id).type_MO,
                            Percent=(count_common + count_general + count_operator) * 100 / cnt
                        )
                        new_rec.save()
                else:
                    count_common = 0
                    count_general = 0
                    count_operator = 0
                    for el_int in intermediate.objects.filter(spec_id__in=list_spec):
                        if el_int.spec_id == spec.id and el_int.level_id == el['level_id']:
                            for el_subint in intermediate.objects.filter(spec_id__in=list_spec):
                                if el_int.cert_id == el_subint.cert_id and el_int.spec_id != el_subint.spec_id \
                                        and el_int.result == el_subint.result:
                                    if el_subint.level_id == ExpLevel.objects.get(name_level='Оператор').id:
                                        count_operator += 1
                                    elif el_subint.level_id == ExpLevel.objects.get(name_level='Первый этап').id:
                                        check_right = True
                                        for el_third in intermediate.objects.filter(spec_id__in=list_spec):
                                            if el_int.cert_id == el_third.cert_id \
                                                    and el_int.level_id != el_third.level_id and \
                                                    el_int.result != el_third.result:
                                                check_right = False
                                        if check_right:
                                            count_common += 1
                                    else:
                                        count_general += 1
                    if res.count != intermediate.objects.filter(spec_id=spec.id).filter(level_id=el['level_id']).count() \
                        or res.coincidence_common != count_common or res.coincidence_general != count_general or res.coincidence_operator != count_operator:
                        cnt = intermediate.objects.filter(spec_id=spec.id).filter(level_id=el['level_id']).count()
                        res.count = intermediate.objects.filter(spec_id=spec.id).filter(level_id=el['level_id']).count()
                        res.coincidence_common = count_common
                        res.coincidence_general = count_general
                        res.coincidence_operator = count_operator
                        res.Percent = (count_common + count_general + count_operator) * 100 / cnt
                        res.save()
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'update_res.html', context)
        pers = periods.objects.all().order_by('-id')
        context = {
            'curr_user': user,
            'curr_group': group,
            'periods': pers
        }
        return render(request, 'tables/choose_update_results.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def add_target(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if group != 'Уполномоченный':
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context) 
        deleg = delegates.objects.filter(email=request.user.email).latest('id')
        if request.method == 'POST':
            mail_spec = request.POST['spec']
            try:
                spec = specialists.objects.filter(period_id=deleg.period_id).filter(email=mail_spec).latest('email')
            except BaseException:
                context = {
                    'title': 'Найдено несколько специалистов с указанным адресом электронной почты или специалист не найден',
                    'curr_user': user,
                    'curr_group': group
                }
                return render(request, 'import/success.html', context)
            atts = request.POST.getlist('atts')
            for el in atts:
                new_rec = targets()
                new_rec.delegate_id = deleg.id
                new_rec.spec_id = spec.id
                cert = certified.objects.get(att_code=el).id
                new_rec.cert_id = cert
                new_rec.save()
            return list_targets_deleg(request)
        targs = targets.objects.filter(delegate_id=deleg.id)
        count=targs.count()
        find = targets_mo.objects.filter(delegate_id=deleg.id)
        if find.count() == 0:
            context = {
                'title': 'Не найдено назначенных МО',
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'import/success.html', context)
        list_mo = []
        for el in find:
            list_mo.append(el.MO_id)
        specs = specialists.objects.filter(period_id=deleg.period_id).filter(MO_id=deleg.MO_id)
        if specs.count() == 0:
            context = {
                'title': 'Не найдено специалистов с вашего МО в текущем периоде аттестации',
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'import/success.html', context)
        list_specs = []
        for el in specs:
            list_specs.append(el.email)
        if count == 0: 
            certs = certified.objects.filter(period_id=deleg.period_id).filter(MO_id__in=list_mo)
            if certs.count() == 0:
                context = {
                    'title': 'Не найдено незначенных аттестуемых',
                    'curr_user': user,
                    'curr_group': group
                }
                return render(request, 'import/success.html', context)
        else:
            list_used = []
            for el in targs:
                list_used.append(el.cert_id)
            certs = certified.objects.filter(period_id=deleg.period_id).filter(MO_id__in=list_mo).exclude(id__in=list_used) 
            if certs.count() == 0:
                context = {
                    'title': 'Не найдено незначенных аттестуемых',
                    'curr_user': user,
                    'curr_group': group
                }
                return render(request, 'import/success.html', context)
        list_certs = []
        for el in certs:
            list_certs.append(el.att_code)
        context = {
            'curr_user': user,
            'curr_group': group,
            'certs': list_certs,
            'specs': list_specs
        }
        return render(request, 'add_target.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def exp_card(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if request.method == 'POST':
            lvls = ExpLevel.objects.all()
            certs = certified.objects.all()
            try:
                spec = specialists.objects.filter(email=request.user.email).latest('id')
            except BaseException:
                context = {
                    'title': "Специалист не найден",
                    'curr_user': user,
                    'curr_group': group
                }
                return render(request, 'import/success.html', context)
            type_mo = MunObr.objects.get(id=spec.MO_id).type_MO
            if group in ('Специалист','Уполномоченный') and type_mo == 'Муниципалитет':
                list_atts=[]
                targs = targets.objects.filter(spec_id=spec.id)
                for el in targs:
                    att_code = certified.objects.get(id=el.cert_id).att_code
                    list_atts.append(att_code)
            check_cert = False
            for el in certs:
                if el.att_code == request.POST['att_code']:
                    check_cert = True
                    id_cert = el.id
            if check_cert == False:
                error = "Аттестуемый не найден"
                if group in ('Специалист','Уполномоченный') and type_mo == 'Муниципалитет':
                    context = {
                        'error': error,
                        'lvls': lvls,
                        'curr_user': user,
                        'atts': list_atts,
                        'curr_group': group
                    }
                else:
                    context = {
                        'error': error,
                        'lvls': lvls,
                        'curr_user': user,
                        'curr_group': group,
                        'gos': 'yes'
                    }
                return render(request, 'main.html', context)
            per = periods.objects.get(id=spec.period_id)
            date_start = per.date_start
            date_end = per.date_end
            today = date.today()
            if date_end < today or today < date_start:
                if type_mo == 'Муниципалитет':
                    context = {
                        'error': 'Период аттестации еще не начался или уже закончился',
                        'lvls': lvls,
                        'curr_user': user,
                        'curr_group': group
                    }
                else:
                    context = {
                        'error': 'Период аттестации еще не начался или уже закончился',
                        'lvls': lvls,
                        'curr_user': user,
                        'curr_group': group,
                        'gos': 'yes'
                    }
                return render(request, 'main.html', context)
            if spec.period_id != certs.get(id=id_cert).period_id:
                if type_mo == 'Муниципалитет':
                    context = {
                        'error': 'Не найден аттестуемый в текущем периоде аттестации',
                        'lvls': lvls,
                        'curr_user': user,
                        'curr_group': group
                     }
                else:
                    context = {
                        'error': 'Не найден аттестуемый в текущем периоде аттестации',
                        'lvls': lvls,
                        'curr_user': user,
                        'curr_group': group,
                        'gos': 'yes'
                    }
                return render(request, 'main.html', context)
            check_double = intermediate.objects.filter(cert_id=id_cert).filter(spec_id=spec.id).count()
            if check_double > 0:
                error = "Экспертная карта на аттестуемого уже добавлена"
                if group in ('Специалист','Уполномоченный') and type_mo == 'Муниципалитет':
                    context = {
                        'error': error,
                        'lvls': lvls,
                        'curr_user': user,
                        'atts': list_atts,
                        'curr_group': group
                    }
                else:
                    context = {
                        'error': error,
                        'lvls': lvls,
                        'curr_user': user,
                        'curr_group': group,
                        'gos': 'yes'
                    }
                return render(request, 'main.html', context)
            new_record = intermediate()
            new_record.spec_id = spec.id
            new_record.cert_id = id_cert
            new_record.level_id = ExpLevel.objects.get(name_level=request.POST['Level_expert']).id
            if request.POST['Level_expert'] == 'Оператор':
                new_record.operator_choice = request.POST['operator_choice']
            else:
                new_record.operator_choice = 'Специалист'
            new_record.result = request.POST['Result']
            new_record.save()
            inter_id = new_record.id
            inter_rec = intermediate.objects.get(id=inter_id)
            spec = specialists.objects.get(id=spec.id)
            att = certified.objects.get(id=id_cert)
            operator_choice = inter_rec.operator_choice
            new_card = ExpCards()
            new_card.period_id = periods.objects.get(id=spec.period_id).id
            new_card.date_add = inter_rec.date_add
            if operator_choice == 'Специалист':
                new_card.type_emp = operator_choice  
            else:
                new_card.type_emp = 'Оператор'
            id_mo_spec = spec.MO_id
            type_mo = MunObr.objects.get(id=id_mo_spec).type_MO
            new_card.type_MO = type_mo
            new_card.MO_expert_id = spec.MO_id
            new_card.FIO_expert = spec.FIO
            new_card.Name_Org_expert = spec.Organization
            new_card.Position_expert_id = spec.Position_id
            new_card.Level_expert_id = inter_rec.level_id
            new_card.MO_att_id = att.MO_id
            new_card.FIO_att = att.FIO
            new_card.Name_Org_att = att.Organization
            new_card.Position_att_id = att.Position_id
            new_card.Category_id = att.Category_id
            new_card.Result = inter_rec.result
            new_card.inter_id = inter_id
            new_card.save()
            return list_cards_spec(request)
        lvls = ExpLevel.objects.all()
        try:
            spec = specialists.objects.filter(email=request.user.email).latest('id')
        except BaseException:
            context = {
            'title': 'Специалист не найден или найдено больше одного специалиста с указанным email',
            'curr_user': user,
            'curr_group': group,
            }
            return render(request, 'import/success.html', context)
        type_mo = MunObr.objects.get(id=spec.MO_id).type_MO
        if type_mo == 'Муниципалитет':
            list_atts = []
            targs = targets.objects.filter(spec_id=spec.id)
            if targs.count() == 0 and group in ('Специалист','Уполномоченный'):
                context = {
                'title': 'Назначений не найдено',
                'curr_user': user,
                'curr_group': group,
                }
                return render(request, 'import/success.html', context)
            for el in targs:
                att_code = certified.objects.get(id=el.cert_id).att_code
                list_atts.append(att_code)
            context = {
                'lvls': lvls,
                'atts': list_atts,
                'curr_user': user,
                'curr_group': group,
            }
        else:
            context = {
                'lvls': lvls,
                'gos': 'yes',
                'curr_user': user,
                'curr_group': group,
            }
        return render(request, 'main.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def list_targets_deleg(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if group != 'Уполномоченный':
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context) 
        try:
            deleg_id = delegates.objects.filter(email=request.user.email).latest('id').id
        except BaseException:
            error = "Уполномоченный не найден или найдено несколько уполномоченных с указанным адресом электронной почты"
            context = {
                'title': error,
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'import/success.html', context)
        targs = targets.objects.filter(delegate_id=deleg_id).order_by('-id')
        if request.method == "POST":
            email = request.POST['value']
            specs = specialists.objects.filter(email__contains=email)
            list_spec = []
            for el in specs:
                list_spec.append(el.id)
            targs = targs.filter(delegate_id=deleg_id).filter(spec_id__in=list_spec)
        if targs.count() == 0:
            context = {
            'curr_user': user,
            'curr_group': group,
            'count': 0
            }
            return render(request, 'targets_deleg.html', context)
        list_targs = []
        voc_targs = {}
        voc_s = {}
        voc_c = {}
        voc_specs = {}
        voc_certs = {}
        for el in targs:
            list_targs.append(el.id)
            id_spec = el.spec_id
            id_cert = el.cert_id
            spec = specialists.objects.get(id=id_spec)
            cert = certified.objects.get(id=id_cert)
            voc_s[el.id] = spec.email
            voc_c[el.id] = cert.att_code
            voc_specs[el.id] = el.spec_id
            voc_certs[el.id] = el.cert_id
        context = {
            'curr_user': user,
            'curr_group': group,
            'targs': list_targs,
            'voc_specs': voc_specs,
            'voc_s': voc_s,
            'voc_c': voc_c,
            'voc_certs': voc_certs,
            
        }
        return render(request, 'targets_deleg.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def list_cards_spec_mo(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        try:
            mail = request.user.email
            deleg = delegates.objects.filter(email=mail).latest('id')
        except BaseException:
            error = "Уполномоченный не найден"
            context = {
                'title': error,
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'import/success.html', context)
        list_spec = []
        mo = MunObr.objects.get(id=deleg.MO_id)
        specs = specialists.objects.filter(MO_id=mo.id)
        for el in specs:
            list_spec.append(el.id)
        count = intermediate.objects.filter(spec_id__in=list_spec).count()
        cards = intermediate.objects.filter(spec_id__in=list_spec).order_by("-date_add")
        lvls = ExpLevel.objects.all()
        voc_specs = {}
        voc_atts = {}
        for el in cards:
            id_spec = el.spec_id
            voc_specs[id_spec] = specialists.objects.get(id=id_spec).email
            id_cert = el.cert_id
            voc_atts[id_cert] = certified.objects.get(id=id_cert).att_code
        context = {
            'cards': cards,
            'lvls': lvls,
            'count': count,
            'voc_atts': voc_atts,
            'voc_specs': voc_specs,
            'curr_user': user,
            'curr_group': group,
        }
        return render(request, 'cards_spec_mo.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def find_list_cards_spec(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        try:
            mail = request.user.email
            deleg = delegates.objects.filter(email=mail).latest('id')
        except BaseException:
            error = "Уполномоченный не найден"
            context = {
                'title': error,
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'import/success.html', context)
        mail = request.POST['mail']
        list_spec = []
        mo = MunObr.objects.get(id=deleg.MO_id)
        specs = specialists.objects.filter(MO_id=mo.id).filter(email__contains=mail)
        for el in specs:
            list_spec.append(el.id)
        count = intermediate.objects.filter(spec_id__in=list_spec).count()
        cards = intermediate.objects.filter(spec_id__in=list_spec).order_by("-date_add")
        lvls = ExpLevel.objects.all()
        voc_specs = {}
        voc_atts = {}
        for el in cards:
            id_spec = el.spec_id
            voc_specs[id_spec] = specialists.objects.get(id=id_spec).email
            id_cert = el.cert_id
            voc_atts[id_cert] = certified.objects.get(id=id_cert).att_code
        context = {
            'cards': cards,
            'lvls': lvls,
            'count': count,
            'voc_atts': voc_atts,
            'voc_specs': voc_specs,
            'curr_user': user,
            'curr_group': group,
        }
        return render(request, 'cards_spec_mo.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def list_cards_spec(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        try:
            spec_id = specialists.objects.filter(email=request.user.email).latest('id').id
        except BaseException:
            error = "Специалист не найден"
            context = {
                'title': error,
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'import/success.html', context)
        count = intermediate.objects.filter(spec_id=spec_id).count()
        cards = intermediate.objects.filter(spec_id=spec_id).order_by("-date_add")
        lvls = ExpLevel.objects.all()
        voc_atts = {}
        for el in cards:
            id = el.cert_id
            voc_atts[id] = certified.objects.filter(id=id).values('att_code').first()
        context = {
            'cards': cards,
            'lvls': lvls,
            'count': count,
            'voc': voc_atts,
            'curr_user': user,
            'curr_group': group,
        }
        return render(request, 'cards_spec.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_target_deleg(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        targets.objects.get(id=id).delete()
        return list_targets_deleg(request)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_spec_card(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        id_user = intermediate.objects.get(id=id).cert_id
        att_code = certified.objects.get(id=id_user).att_code
        info = 'Карта с аттестуемым '+att_code+' успешно удалена'
        intermediate.objects.get(id=id).delete()
        mail = request.user.email
        try:
            spec_id = specialists.objects.get(email=mail)
        except BaseException:
            spec_id = specialists.objects.filter(email=mail).latest('id')
        cards = intermediate.objects.filter(spec_id=spec_id.id)
        count = cards.count()
        lvls = ExpLevel.objects.all()
        voc_atts = {}
        for el in cards:
            id = el.cert_id
            voc_atts[id] = certified.objects.filter(id=id).values('att_code').first()
        context = {
            'curr_user': user,
            'curr_group': group,
            'cards': cards,
            'lvls': lvls,
            'count': count,
            'voc': voc_atts,
            'info': info
        }
        return render(request, 'cards_spec.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def change_password_delegate(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        deleg = delegates.objects.get(id=id)
        if request.method == 'POST':
            if request.POST['pass'] != request.POST['pass2']:
                context = {
                    'curr_user': user,
                    'curr_group': group,
                    'error': 'Введенные пароли не совпадают',
                    'spec': deleg
                }
                return render(request, 'registration/change_pass.html', context)
            mail = deleg.email
            spec_user = User.objects.get(email=mail)
            spec_user.set_password(request.POST['pass'])
            spec_user.save()
            context = {
                'curr_user': user,
                'curr_group': group,
                'title': 'Пароль был успешно изменен'
            }
            return render(request, 'import/success.html', context)
        else:
            context = {
                'curr_user': user,
                'curr_group': group,
                'spec': deleg
            }
            return render(request, 'registration/change_pass.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def change_password(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        spec = specialists.objects.get(id=id)
        if request.method == 'POST':
            if request.POST['pass'] != request.POST['pass2']:
                context = {
                    'curr_user': user,
                    'curr_group': group,
                    'error': 'Введенные пароли не совпадают',
                    'spec': spec
                }
                return render(request, 'registration/change_pass.html', context)
            mail = spec.email
            spec_user = User.objects.get(email=mail)
            spec_user.set_password(request.POST['pass'])
            spec_user.save()
            context = {
                'curr_user': user,
                'curr_group': group,
                'title': 'Пароль был успешно изменен'
            }
            return render(request, 'import/success.html', context)
        else:
            context = {
                'curr_user': user,
                'curr_group': group,
                'spec': spec
            }
            return render(request, 'registration/change_pass.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def reset_pass(request):
    return render(request, 'registration/password_reset.html')


def login_user(request):
    username = request.POST['username']
    password = request.POST['password']
    user = authenticate(request, username=username, password=password)
    if user is not None:
        if check_network(request) == False:
            list_group = user.groups.values_list('name', flat=True)
            group = list_group[0]
            if group == 'Специалист' or group == 'Уполномоченный':
                mail = User.objects.get(username=username).email
                per = periods.objects.latest('id')
                specs = specialists.objects.filter(period_id=per.id)
                delegs = delegates.objects.filter(period_id=per.id)
                check = False
                for el in specs:
                    if mail == el.email:
                        check = True
                        break
                for el in delegs:
                    if mail == el.email:
                        check = True
                        break
                if check == False:
                    return HttpResponse("Специалист/уполномоченный не найден на период "+per.name_period)
        login(request, user)
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/">')
    else:
        info = 'Неудачная попытка аутентификации'
        return render(request, 'registration/login.html', {'info': info})


def registr(request):
    if request.method == 'POST':
        logins = User.objects.all()
        for el in logins:
            if el.username == request.POST['username']:
                error = "Пользователь с таким логином уже существует"
                return render(request, 'registration/registration.html', { 'error': error})
            if el.email == request.POST['email']:
                error = "Пользователь с таким email уже существует"
                return render(request, 'registration/registration.html', { 'error': error})
        specs = specialists.objects.all()
        delegs = delegates.objects.all()
        count = 0
        for el in specs:
            if el.email == request.POST['email']:
                count += 1
        if count == 0:
            for el in delegs:
                if el.email == request.POST['email']:
                    count += 1
        if count == 0:
            error = "Специалист/уполномоченный с таким email не найден"
            return render(request, 'registration/registration.html', {'error': error})
        if request.POST['password'] != request.POST['password2']:
            error = "Введенные пароли не совпадают"
            return render(request, 'registration/registration.html', {'error': error})
        new_user = User()
        new_user.username = request.POST['username']
        new_user.email = request.POST['email']
        new_user.set_password(request.POST['password'])
        new_user.save()
        new_user.refresh_from_db()
        check_spec = True
        for el in delegs:
             if el.email == request.POST['email']:
                check_spec = False
        if check_spec:
            my_group = Group.objects.get(name='Специалист')
        else:
            my_group = Group.objects.get(name='Уполномоченный')
        my_group.user_set.add(new_user)
        return HttpResponse('<center><h2>Регистрация успешно завершена</h2><br>'
                            'Переход на страницу входа через 3 секунды</center>'
                            '<meta http-equiv="refresh" content="3; URL=/accounts/login/">')
    else:
        return render(request, 'registration/registration.html')


def start(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        context = {
            'curr_user': user,
            'curr_group': group
        }
        return render(request, 'start.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')
        
 
def choose_expcards(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            type_mo = request.POST['type_mo']
            type_emp = request.POST['type_emp']
            period = request.POST.get('period')
            return all_expcards(request, period, type_mo, type_emp)
        else:
            pers = periods.objects.all().order_by('-id')
            count = pers.count()
            context = {
                'curr_user': user,
                'curr_group': group,
                'count': count,
                'periods': pers
            }
            return render(request, 'tables/choose_cards.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def all_expcards(request, period, type_mo, type_emp):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        typ_mo = uri_to_iri(type_mo)
        typ_emp = uri_to_iri(type_emp)
        cards = ExpCards.objects.filter(period_id=period).filter(type_MO=typ_mo).filter(type_emp=typ_emp).order_by('-id')
        mo_exp = cards.order_by("MO_expert").values('MO_expert').distinct()
        pos_exp = cards.order_by("Position_expert").values('Position_expert').distinct()
        mo_att = cards.order_by("MO_att").values('MO_att').distinct()
        pos_att = cards.order_by("Position_att").values('Position_att').distinct()
        count = cards.count()
        try:
            ord_by = request.GET.get('order_by')
            cards = cards.order_by(ord_by)
        except BaseException:
            pass
        count = 0
        try:
            filter_mo_exp = request.POST.getlist('mo_exp')
            posts = cards.filter(MO_expert__in=filter_mo_exp)
            count = posts.cards()
        except BaseException:
            pass
        try:
            filter_mo_att = request.POST.getlist('mo_att')
            posts = cards.filter(MO_expert__in=filter_mo_att)
            count = posts.cards()
        except BaseException:
            pass
        if count == 0:
            paginator = Paginator(cards, 50)
            try:
                page = request.GET.get('page')
            except BaseException:
                page = 1
            try:
                posts = paginator.page(page)
            except PageNotAnInteger:
                # Если страница не является целым числом, поставим первую страницу
                posts = paginator.page(1)
            except EmptyPage:
                # Если страница больше максимальной, доставить последнюю страницу результатов
                posts = paginator.page(paginator.num_pages)
            count = cards.count()
        list_mo_exp = []
        for el in mo_exp:
            name = MunObr.objects.get(id=el['MO_expert']).name_MO
            list_mo_exp.append(name)
        list_pos_exp = []
        for el in pos_exp:
            name = Position.objects.get(id=el['Position_expert']).name_pos
            list_pos_exp.append(name)
        list_mo_att = []
        for el in mo_att:
            name = MunObr.objects.get(id=el['MO_att']).name_MO
            list_mo_att.append(name)
        list_pos_att = []
        for el in pos_att:
            name = Position.objects.get(id=el['Position_att']).name_pos
            list_pos_att.append(name)
        per = periods.objects.get(id=period)
        context = {
            'filter': 'yes',
            'curr_user': user,
            'curr_group': group,
            'mo': typ_mo,
            'title': 'yes',
            'period': per,
            'mo_exp': list_mo_exp,
            'pos_exp': list_pos_exp,
            'mo_att': list_mo_att,
            'pos_att': list_pos_att,
            'emp': typ_emp,
            'cards': posts,
            'count': count,
            'order_by': ord_by
        }
        return render(request, 'tables/all_expcards.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def filter_expcards(request, period, type_mo, type_emp):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        typ_mo = uri_to_iri(type_mo)
        typ_emp = uri_to_iri(type_emp)
        tabl = ExpCards.objects.filter(period_id=period).filter(type_MO=typ_mo).filter(type_emp=typ_emp)
        mo_exp = tabl.values('MO_expert_id').distinct()
        pos_exp = tabl.values('Position_expert_id').distinct()
        mo_att = tabl.values('MO_att_id').distinct()
        pos_att = tabl.values('Position_att_id').distinct()
        filter_mo_exp = []
        filter_pos_exp = []
        filter_mo_att = []
        filter_pos_att = []
        list_filter_mo_exp = []
        list_filter_pos_exp = []
        list_filter_mo_att = []
        list_filter_pos_att = []
        if len(request.POST.getlist('filter_mo_exp')) > 0:
            filter_mo = request.POST.getlist('filter_mo_exp')
            for el in filter_mo:
                list_filter_mo_exp.append(el)
            tabl = tabl.filter(MO_expert_id__in=list_filter_mo_exp)
            pos_exp = tabl.values('Position_expert_id').distinct()
            mo_att = tabl.values('MO_att_id').distinct()
            pos_att = tabl.values('Position_att_id').distinct()
        if len(request.POST.getlist('filter_pos_exp')) > 0:
            filter_pos = request.POST.getlist('filter_pos_exp')
            for el in filter_pos:
                list_filter_pos_exp.append(el)
            tabl = tabl.filter(Position_expert_id__in=list_filter_pos_exp)
            mo_exp = tabl.values('MO_expert_id').distinct()
            mo_att = tabl.values('MO_att_id').distinct()
            pos_att = tabl.values('Position_att_id').distinct()
        if len(request.POST.getlist('filter_mo_att')) > 0:
            filter_mo = request.POST.getlist('filter_mo_att')
            for el in filter_mo:
                list_filter_mo_att.append(el)
            tabl = tabl.filter(MO_att_id__in=list_filter_mo_att)
            mo_exp = tabl.values('MO_expert_id').distinct()
            pos_exp = tabl.values('Position_expert_id').distinct()
            pos_att = tabl.values('Position_att_id').distinct()
        if len(request.POST.getlist('filter_pos_att')) > 0:
            filter_pos = request.POST.getlist('filter_pos_att')
            for el in filter_pos:
                list_filter_pos_att.append(el)
            tabl = tabl.filter(Position_att_id__in=list_filter_pos_att)
            mo_exp = tabl.values('MO_expert_id').distinct()
            pos_exp = tabl.values('Position_expert_id').distinct()
            mo_att = tabl.values('MO_att_id').distinct()
        if len(request.POST.getlist('mo_exp'))!=0:
            filter_mo_exp = request.POST.getlist('mo_exp')
            list_mo_exp = []
            for el in filter_mo_exp:
                id_filter = MunObr.objects.get(name_MO=el).id
                list_mo_exp.append(id_filter)
            filter_mo_exp = list_mo_exp
            posts = tabl.filter(MO_expert_id__in=list_mo_exp).order_by("FIO_expert")
            pos_exp = posts.order_by("Position_expert_id").values('Position_expert_id').distinct()
            mo_att = posts.order_by("MO_att_id").values('MO_att_id').distinct()
            pos_att = posts.order_by("Position_att_id").values('Position_att_id').distinct()
            count = posts.count()
        if len(request.POST.getlist('pos_exp'))!=0:
            filter_pos_exp = request.POST.getlist('pos_exp')
            list_pos_exp = []
            for el in filter_pos_exp:
                id_filter = Position.objects.filter(type_pos='Специалист').get(name_pos=el).id
                list_pos_exp.append(id_filter)
            filter_pos_exp = list_pos_exp
            posts = tabl.filter(Position_expert_id__in=list_pos_exp).order_by("FIO_expert")
            mo_exp = posts.order_by("MO_expert_id").values('MO_expert_id').distinct()
            mo_att = posts.order_by("MO_att_id").values('MO_att_id').distinct()
            pos_att = posts.order_by("Position_att_id").values('Position_att_id').distinct()
            count = posts.count()
        if len(request.POST.getlist('mo_att'))!=0:
            filter_mo_att = request.POST.getlist('mo_att')
            list_mo_att = []
            for el in filter_mo_att:
                id_filter = MunObr.objects.get(name_MO=el).id
                list_mo_att.append(id_filter)
            filter_mo_att = list_mo_att
            posts = tabl.filter(MO_att_id__in=list_mo_att).order_by("FIO_att")
            mo_exp = posts.order_by("MO_expert_id").values('MO_expert_id').distinct()
            pos_exp = posts.order_by("Position_expert_id").values('Position_expert_id').distinct()
            pos_att = posts.order_by("Position_att_id").values('Position_att_id').distinct()
            count = posts.count()
        if len(request.POST.getlist('pos_att'))!=0:
            filter_mo_exp = request.POST.getlist('pos_att')
            list_pos_att = []
            for el in filter_mo_exp:
                id_filter = Position.objects.filter(type_pos='Аттестуемый').get(name_pos=el).id
                list_pos_att.append(id_filter)
            filter_pos_att = list_pos_att
            posts = tabl.filter(Position_att_id__in=list_pos_att).order_by("FIO_att")
            mo_exp = posts.order_by("MO_expert_id").values('MO_expert_id').distinct()
            pos_exp = posts.order_by("Position_expert_id").values('Position_expert_id').distinct()
            mo_att = posts.order_by("MO_att_id").values('MO_att_id').distinct()
            count = posts.count()
        try:
            ord_by = request.GET.get('order_by')
            posts = posts.order_by(ord_by)
        except BaseException:
            pass
        list_mo_exp = []
        for el in mo_exp:
            name = MunObr.objects.get(id=el['MO_expert_id']).name_MO
            list_mo_exp.append(name)
        list_mo_att = []
        for el in mo_att:
            name = MunObr.objects.get(id=el['MO_att_id']).name_MO
            list_mo_att.append(name)
        list_pos_exp = []
        for el in pos_exp:
            name = Position.objects.filter(type_pos='Специалист').get(id=el['Position_expert_id']).name_pos
            list_pos_exp.append(name)
        list_pos_att = []
        for el in pos_att:
            name = Position.objects.filter(type_pos='Аттестуемый').get(id=el['Position_att_id']).name_pos
            list_pos_att.append(name)
        context = {
            'curr_user': user,
            'curr_group': group,
            'mo': typ_mo,
            'emp': typ_emp,
            'title': 'yes',
            'mo_exp': list_mo_exp,
            'pos_exp': list_pos_exp,
            'mo_att': list_mo_att,
            'pos_att': list_pos_att,
            'filter_mo_exp': filter_mo_exp,
            'filter_pos_exp': filter_pos_exp,
            'filter_mo_att': filter_mo_att,
            'filter_pos_att': filter_pos_att,
            'cards': posts,
            'count': count,
            'order_by': ord_by
        }
        return render(request, 'tables/all_expcards.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def filter_specs(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        per = periods.objects.get(id=period)
        mo = specialists.objects.filter(period_id=period).order_by('MO_id').values('MO_id').distinct()
        pos = specialists.objects.filter(period_id=period).order_by('Position_id').values('Position_id').distinct()
        specs = specialists.objects.filter(period_id=period)
        list_filter_mo = []
        list_filter_pos = []
        if len(request.POST.getlist('filter_mo')) > 0:
            filter_mo = request.POST.getlist('filter_mo')
            for el in filter_mo:
                list_filter_mo.append(el)
            specs = specs.filter(MO_id__in=list_filter_mo)
            pos = specs.order_by('Position_id').values('Position_id').distinct()
        if len(request.POST.getlist('filter_pos')) > 0:
            filter_pos = request.POST.getlist('filter_pos')
            for el in filter_pos:
                list_filter_pos.append(el)
            specs = specs.filter(Position_id__in=list_filter_pos)
            mo = specs.order_by('MO_id').values('MO_id').distinct()
        if len(request.POST.getlist('mo')) != 0:
            mo_filt = request.POST.getlist('mo')
            list_mo = []
            for el in mo_filt:
                name = MunObr.objects.get(name_MO=el).id
                list_mo.append(name)
            specs = specs.filter(MO_id__in=list_mo).order_by("FIO")
            list_filter_mo = list_mo
            pos = specs.order_by('Position_id').values('Position_id').distinct()
        if len(request.POST.getlist('pos')) != 0:
            pos_filt = request.POST.getlist('pos')
            list_pos = []
            for el in pos_filt:
                name = Position.objects.filter(type_pos='Специалист').get(name_pos=el).id
                list_pos.append(name)
            specs = specs.filter(Position_id__in=list_pos).order_by("FIO")
            list_filter_pos = list_pos
            mo = specs.order_by('MO_id').values('MO_id').distinct()
        try:
            count = specs.count()
        except BaseException:
            pass
        list_mo = []
        for el in mo:
            name = MunObr.objects.get(id=el['MO_id']).name_MO
            list_mo.append(name)
        list_pos = []
        for el in pos:
            name = Position.objects.get(id=el['Position_id']).name_pos
            list_pos.append(name) 
        context = {
            'curr_user': user,
            'period': per,
            'curr_group': group,
            'filter': 'yes',
            'mo': list_mo,
            'pos': list_pos,
            'filter_mo': list_filter_mo,
            'filter_pos': list_filter_pos,
            'specs': specs,
            'count': count,
        }
        return render(request, 'import/specs_list.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def filter_delegates(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        per = periods.objects.get(id=period)
        mo = delegates.objects.filter(period_id=period).order_by('MO_id').values('MO_id').distinct()
        delegs = delegates.objects.filter(period_id=period)
        list_filter_mo = []
        if len(request.POST.getlist('filter_mo')) > 0:
            filter_mo = request.POST.getlist('filter_mo')
            for el in filter_mo:
                list_filter_mo.append(el)
            delegs = delegs.filter(MO_id__in=list_filter_mo)
        if len(request.POST.getlist('mo')) != 0:
            mo_filt = request.POST.getlist('mo')
            list_mo = []
            for el in mo_filt:
                name = MunObr.objects.get(name_MO=el).id
                list_mo.append(name)
            delegs = delegs.filter(MO_id__in=list_mo).order_by("FIO")
            list_filter_mo = list_mo
        try:
            count = delegs.count()
        except BaseException:
            pass
        list_mo = []
        for el in mo:
            name = MunObr.objects.get(id=el['MO_id']).name_MO
            list_mo.append(name)
        context = {
            'curr_user': user,
            'period': per,
            'curr_group': group,
            'filter': 'yes',
            'mo': list_mo,
            'filter_mo': list_filter_mo,
            'delegates': delegs,
            'count': count,
        }
        return render(request, 'import/delegates_list.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def filter_atts(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        mo = certified.objects.filter(period_id=period).order_by('MO_id').values('MO_id').distinct()
        pos = certified.objects.filter(period_id=period).order_by('Position_id').values('Position_id').distinct()
        cert = certified.objects.filter(period_id=period)
        list_filter_mo = []
        list_filter_pos = []
        if len(request.POST.getlist('filter_mo')) > 0:
            filter_mo = request.POST.getlist('filter_mo')
            for el in filter_mo:
                list_filter_mo.append(el)
            cert = cert.filter(MO_id__in=list_filter_mo)
            pos = cert.order_by('Position_id').values('Position_id').distinct()
        if len(request.POST.getlist('filter_pos')) > 0:
            filter_pos = request.POST.getlist('filter_pos')
            for el in filter_pos:
                list_filter_pos.append(el)
            cert = cert.filter(Position_id__in=list_filter_pos)
            mo = cert.order_by('MO_id').values('MO_id').distinct()
        if len(request.POST.getlist('mo')) != 0:
            mo_filt = request.POST.getlist('mo')
            list_mo = []
            for el in mo_filt:
                name = MunObr.objects.get(name_MO=el).id
                list_mo.append(name)
            cert = cert.filter(MO_id__in=list_mo).order_by("FIO")
            list_filter_mo = list_mo
            pos = cert.order_by('Position_id').values('Position_id').distinct()
        if len(request.POST.getlist('pos')) != 0:
            pos_filt = request.POST.getlist('pos')
            list_pos = []
            for el in pos_filt:
                name = Position.objects.filter(type_pos='Аттестуемый').get(name_pos=el).id
                list_pos.append(name)
            cert = cert.filter(Position_id__in=list_pos).order_by("FIO")
            list_filter_pos = list_pos
            mo = cert.order_by('MO_id').values('MO_id').distinct()
        try:
            count = cert.count()
        except BaseException:
            pass
        list_mo = []
        for el in mo:
            name = MunObr.objects.get(id=el['MO_id']).name_MO
            list_mo.append(name)
        list_pos = []
        for el in pos:
            name = Position.objects.get(id=el['Position_id']).name_pos
            list_pos.append(name)    
        per = periods.objects.get(id=period)
        context = {
            'curr_user': user,
            'curr_group': group,
            'period': per,
            'filter': 'yes',
            'mo': list_mo,
            'pos': list_pos,
            'filter_mo': list_filter_mo,
            'filter_pos': list_filter_pos,
            'cert': cert,
            'count': count,
        }
        return render(request, 'import/cert_list.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def choose_criterias(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            per = request.POST.get('period')
            return all_criterias(request, per)
        pers = periods.objects.all().order_by('-id')
        count = pers.count()
        context = {
            'curr_user': user,
            'curr_group': group,
            'count': count,
            'periods': pers
        }
        return render(request, 'tables/choose_criterias.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def all_criterias(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        criterias = criteria_export.objects.filter(period_id=period).order_by('-id')
        count = criterias.count()
        try:
            ord_by = request.GET.get('order_by')
            criterias = criterias.order_by(ord_by)
        except BaseException:
            pass
        paginator = Paginator(criterias, 50)
        try:
            page = request.GET.get('page')
        except BaseException:
            page = 1
        try:
            posts = paginator.page(page)
        except PageNotAnInteger:
            # Если страница не является целым числом, поставим первую страницу
            posts = paginator.page(1)
        except EmptyPage:
            # Если страница больше максимальной, доставить последнюю страницу результатов
            posts = paginator.page(paginator.num_pages)
        per = periods.objects.get(id=period)
        context = {
            'curr_user': user,
            'curr_group': group,
            'criterias': posts,
            'count': count,
            'period': per,
            'order_by': ord_by
        }
        return render(request, 'tables/all_criterias.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def cert_results_from_targets(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        interm_recs = intermediate.objects.filter(cert_id=id)
        if interm_recs.count() == 0:
            context = {
            'curr_user': user,
            'curr_group': group,
            'count': 0
            }
            return render(request, 'tables/all_expcards.html', context)
        query = ExpCards.objects.filter(inter_id__in=interm_recs)
        for el in query:
            type_mo = el.type_MO
            type_emp = el.type_emp
        count = query.count()
        fio = certified.objects.get(id=id).FIO
        title = "Карты аттестуемого: "+fio
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': title,
            'cards': query,
            'mo': type_mo,
            'emp': type_emp,
            'count': count,
            'cert_id': id,
        }
        return render(request, 'tables/all_expcards.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def cert_results_from_expcard(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        interm = intermediate.objects.get(id=id).cert_id
        interm_recs = intermediate.objects.filter(cert_id=interm)
        if interm_recs.count() == 0:
            context = {
            'curr_user': user,
            'curr_group': group,
            'count': 0
            }
            return render(request, 'tables/all_expcards.html', context)
        query = ExpCards.objects.filter(inter_id__in=interm_recs)
        count = query.count()
        for el in query:
            type_mo = el.type_MO
            type_emp = el.type_emp
        fio = certified.objects.get(id=interm).FIO
        title = "Карты аттестуемого: "+fio
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': title,
            'cards': query,
            'type_mo': type_mo,
            'type_emp': type_emp,
            'count': count,
            'cert_id': id,
        }
        return render(request, 'tables/all_expcards.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def cert_results(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        interm_recs = intermediate.objects.filter(cert_id=id)
        if interm_recs.count() == 0:
            context = {
            'curr_user': user,
            'curr_group': group,
            'count': 0
            }
            return render(request, 'tables/all_expcards.html', context)
        query = ExpCards.objects.filter(inter_id__in=interm_recs)
        for el in query:
            type_mo = el.type_MO
            type_emp = el.type_emp
        count = query.count()
        fio = certified.objects.get(id=id).FIO
        title = "Карты аттестуемого: "+fio
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': title,
            'cards': query,
            'inters': interm_recs,
            'mo': type_mo,
            'emp': type_emp,
            'count': count,
            'cert_id': id,
        }
        return render(request, 'tables/all_expcards.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def summary_results(request, fio, period, mo, pos):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        fio_uri = uri_to_iri(fio)
        fio_iri = ' '.join(fio_uri.split('%20'))
        id_spec = specialists.objects.filter(FIO=fio_iri).filter(MO_id=mo).filter(period_id=period).get(Position_id=pos).id
        return spec_results(request, id_spec)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def spec_results(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        interm_recs = intermediate.objects.filter(spec_id=id)
        if interm_recs.count() == 0:
            context = {
            'curr_user': user,
            'curr_group': group,
            'count': 0
            }
            return render(request, 'tables/all_expcards.html', context)
        query = ExpCards.objects.filter(inter_id__in=interm_recs)
        for el in query:
            type_mo = el.type_MO
            type_emp = el.type_emp
        count = query.count()
        fio = specialists.objects.get(id=id).FIO
        title = "Карты аттестуемого: "+fio
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': title,
            'cards': query,
            'mo': type_mo,
            'emp': type_emp,
            'count': count
        }
        return render(request, 'tables/all_expcards.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')

def choose_results(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            typ = request.POST['type']
            per = request.POST.get('period')
            return all_results(request, typ, per, 1)
        else:
            pers = periods.objects.all().order_by('-id')
            count = pers.count()
            context = {
                'curr_user': user,
                'curr_group': group,
                'count': count,
                'periods': pers
            }
            return render(request, 'tables/choose_results.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def all_results(request, typ, period, page):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        type_iri = uri_to_iri(typ)
        results = Summary_table.objects.filter(period_id=period).filter(type_MO=type_iri)
        try:
            ord_by = request.GET.get('order_by')
            results = results.order_by(ord_by)
        except BaseException:
            pass
        count = results.count()
        paginator = Paginator(results, 50)
        try:
            posts = paginator.page(page)
        except PageNotAnInteger:
            # Если страница не является целым числом, поставим первую страницу
            posts = paginator.page(1)
        except EmptyPage:
            # Если страница больше максимальной, доставить последнюю страницу результатов
            posts = paginator.page(paginator.num_pages)
        title = type_iri
        per = periods.objects.get(id=period)
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': title,
            'results': posts,
            'count': count,
            'period': per,
            'order_by': ord_by
        }
        return render(request, 'tables/all_results.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def import_excel(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            period = request.POST['period']
            typ = request.POST['type']
            if typ == 'Аттестуемые': 
                today = str(date.today())
                count = certified.objects.filter(period_id=period).count()
                if count == 0:
                   att_code = today
                   newmax = 1
                else:
                    id_max = certified.objects.all().aggregate(max=Max('id')).get('max')
                    code = certified.objects.get(id=id_max).att_code
                    newmax = int(code[8:])
                    newmax += randint(0, 50)
                    att_code = today[:8] + str(newmax)
                file = xl.readxl(request.FILES['excel'])
                rows = file.ws(ws='Лист1').rows
                i = 3
                error = []
                for row in islice(rows, 2, None):
                    error_mo = ""
                    error_pos = ""
                    error_cat = ""
                    error_add = ""
                    error_attform = ""
                    exfio = file.ws(ws='Лист1').index(row=i, col=2)
                    cor_str = exfio.strip()
                    fio = ' '.join(cor_str.split())
                    mo = str(file.ws(ws='Лист1').index(row=i, col=3)).strip()
                    pos = str(file.ws(ws='Лист1').index(row=i, col=5)).strip()
                    category = str(file.ws(ws='Лист1').index(row=i, col=7)).strip()
                    attform = str(file.ws(ws='Лист1').index(row=i, col=9)).strip()
                    try:
                        id_mo = MunObr.objects.get(name_MO=mo).id
                    except BaseException:
                        error_mo = fio + ": не найдено МО '" + mo + "'"
                    try:
                        id_pos = Position.objects.filter(type_pos='Аттестуемый').get(name_pos=pos).id
                    except BaseException:
                        error_pos = fio + ": не найдена должность '" + pos + "'"
                    try:
                        id_cat = AttCategories.objects.get(type_cat=category).id
                    except BaseException:
                        error_cat = fio + ": не найдена категория '" + category + "'"
                    try:
                        id_attform = AttForm.objects.get(name=attform).id
                    except BaseException:
                        error_attform = fio + ": не найдена форма аттестации '" + attform + "'"
                    if len(error_mo) == 0 and len(error_pos) == 0 and len(error_cat) == 0 and len(error_attform) == 0:
                        new_att = certified()
                        new_att.period_id = period
                        new_att.att_code = att_code
                        new_att.MO_id = id_mo
                        new_att.FIO = fio
                        new_att.Organization = file.ws(ws='Лист1').index(row=i, col=4)
                        new_att.Position_id = id_pos
                        new_att.Category_id = id_cat
                        new_att.att_form_id = id_attform
                        new_att.target = "Нет"
                        try:
                            new_att.save()
                        except BaseException as e:
                            error_add = fio + ": не удалось загрузить аттестуемого," + str(e)
                    newmax += randint(10, 50)
                    att_code = today[:8] + str(newmax)
                    i += 1
                    if len(error_mo) != 0:
                        error.append(error_mo)
                    if len(error_pos) != 0:
                        error.append(error_pos)
                    if len(error_cat) != 0:
                        error.append(error_cat)
                    if len(error_attform) != 0:
                        error.append(error_attform)
                    if len(error_add) != 0:
                        error.append(error_add)
                cert = certified.objects.filter(period_id=period)
                count = cert.count()
                try:
                    per = periods.objects.get(id=period)
                except BaseException:
                    context = {
                    'curr_user': user,
                    'curr_group': group,
                    'title': 'Период не найден или найдено более двух периодов'
                    }
                    return render(request, 'import/success.html', context)
                context = {
                    'curr_user': user,
                    'curr_group': group,
                    'error': error,
                    'cert': cert,
                    'count': count,
                    'period': per
                }
                return render(request, 'import/cert_list.html', context)
            elif typ == 'Специалисты':
                file = xl.readxl(request.FILES['excel'])
                worksheets = file.ws_names
                error = []
                for sheet in worksheets:
                    rows = file.ws(ws=sheet).rows
                    i = 3
                    error_sheet = ""
                    for row in islice(rows, 2, None):
                        error_mo = ""
                        error_pos = ""
                        error_add = ""
                        exfio = str(file.ws(ws=sheet).index(row=i, col=2))
                        cor_str = exfio.strip()
                        fio = ' '.join(cor_str.split())
                        mo = str(file.ws(ws=sheet).index(row=i, col=3)).strip()
                        pos = str(file.ws(ws=sheet).index(row=i, col=5)).strip()
                        try:
                            id_mo = MunObr.objects.get(name_MO=mo).id
                        except BaseException:
                            error_mo = fio + ": не найдено МО '" + mo + "' на листе " + sheet
                        try:
                            id_pos = Position.objects.filter(type_pos='Специалист').get(name_pos=pos).id
                        except BaseException:
                            error_pos = fio + ": не найдена должность '" + pos + "' на листе " + sheet
                        if len(error_mo) == 0 and len(error_pos) == 0:
                            new_spec = specialists()
                            new_spec.period_id = period
                            new_spec.MO_id = id_mo
                            new_spec.FIO = fio
                            new_spec.Organization = str(file.ws(ws=sheet).index(row=i, col=4)).strip()
                            new_spec.Position_id = id_pos
                            new_spec.email = str(file.ws(ws=sheet).index(row=i, col=9)).strip()
                            try:
                                new_spec.save()
                            except BaseException:
                                error_add = "Ошибка при попытке загрузить специалиста " + fio
                        i += 1
                        if len(error_mo) != 0:
                            error.append(error_mo)
                        if len(error_pos) != 0:
                            error.append(error_pos)
                        if len(error_add) != 0:
                            error.append(error_add)
                specs = specialists.objects.filter(period_id=period)
                count = specs.count()
                try:
                    per = periods.objects.get(id=period)
                except BaseException:
                    context = {
                    'curr_user': user,
                    'curr_group': group,
                    'title': 'Период не найден или найдено более двух периодов'
                    }
                    return render(request, 'import/success.html', context)
                context = {
                    'curr_user': user,
                    'error': error,
                    'curr_group': group,
                    'period': per,
                    'specs': specs,
                    'count': count,
                }
                return render(request, 'import/specs_list.html', context)
            else:
                file = xl.readxl(request.FILES['excel'])
                worksheets = file.ws_names
                error = []
                for sheet in worksheets:
                    rows = file.ws(ws=sheet).rows
                    i = 3
                    error_sheet = ""
                    for row in islice(rows, 2, None):
                        error_mo = ""
                        error_add = ""
                        exfio = str(file.ws(ws=sheet).index(row=i, col=3))
                        cor_str = exfio.strip()
                        fio = ' '.join(cor_str.split())
                        mo = str(file.ws(ws=sheet).index(row=i, col=2)).strip()
                        try:
                            id_mo = MunObr.objects.get(name_MO=mo).id
                        except BaseException:
                            error_mo = fio + ": не найдено МО '" + mo + "' на листе " + sheet
                        if len(error_mo)==0:
                            new_spec = delegates()
                            new_spec.period_id = period
                            new_spec.MO_id = id_mo
                            new_spec.FIO = fio
                            new_spec.Organization = str(file.ws(ws=sheet).index(row=i, col=5)).strip()
                            new_spec.email = str(file.ws(ws=sheet).index(row=i, col=7)).strip()
                            try:
                                new_spec.save()
                            except BaseException:
                                error_add = "Ошибка при попытке загрузить уполномоченного " + fio
                        i += 1
                        if len(error_mo) != 0:
                            error.append(error_mo)
                        if len(error_add) != 0:
                            error.append(error_add)
                delegs = delegates.objects.filter(period_id=period)
                count = delegs.count()
                try:
                    per = periods.objects.get(id=period)
                except BaseException:
                    context = {
                    'curr_user': user,
                    'curr_group': group,
                    'title': 'Период не найден или найдено более двух периодов'
                    }
                    return render(request, 'import/success.html', context)
                context = {
                    'curr_user': user,
                    'error': error,
                    'curr_group': group,
                    'period': per,
                    'delegates': delegs,
                    'count': count,
                }
                return render(request, 'import/delegates_list.html', context)
        else:
            pers = periods.objects.all().order_by('-id')
            if pers.count() == 0:
                context = {
                'curr_user': user,
                'curr_group': group,
                'title': 'Периоды не найдены'
                }
                return render(request, 'import/success.html', context)
            context = {
                'curr_user': user,
                'curr_group': group,
                'periods': pers
            }
            return render(request, 'import/import.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def edit_cert(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            cert = certified.objects.get(id=id)
            cert.MO_id = MunObr.objects.get(name_MO=request.POST['mo']).id
            cert.FIO = request.POST['FIO']
            cert.Organization = request.POST['Org']
            cert.Position_id = Position.objects.filter(type_pos='Аттестуемый').get(name_pos=request.POST['pos']).id
            cert.Category_id = AttCategories.objects.get(type_cat=request.POST['cat']).id
            cert.att_form_id = AttForm.objects.get(name=request.POST['attform']).id
            cert.save()
            context = {
                'curr_user': user,
                'curr_group': group,
                'success': 'Данные успешно изменены'
            }
            return render(request, 'start.html', context)
        else:
            cert = certified.objects.get(id=id)
            mo = MunObr.objects.all()
            pos = Position.objects.filter(type_pos='Аттестуемый')
            cat = AttCategories.objects.all()
            att_form = AttForm.objects.all()
            context = {
                'cert': cert,
                'mo': mo,
                'pos': pos,
                'cat': cat,
                'att_form': att_form,
                'curr_user': user,
                'curr_group': group
            }
            return render (request, 'edit_cert.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def edit_delegate(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            deleg = delegates.objects.get(id=id)
            deleg.MO_id = MunObr.objects.get(name_MO=request.POST['mo']).id
            deleg.FIO = request.POST['FIO']
            deleg.Organization = request.POST['Org']
            deleg.email = request.POST['email']
            deleg.save()
            context = {
                'curr_user': user,
                'curr_group': group,
                'success': 'Данные успешно изменены'
            }
            return render(request, 'start.html', context)
        else:
            deleg = delegates.objects.get(id=id)
            mo = MunObr.objects.all()
            context = {
                'delegate': deleg,
                'mo': mo,
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'edit_delegate.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def edit_spec(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            spec = specialists.objects.get(id=id)
            spec.MO_id = MunObr.objects.get(name_MO=request.POST['mo']).id
            spec.FIO = request.POST['FIO']
            spec.Organization = request.POST['Org']
            spec.Position_id = Position.objects.filter(type_pos='Специалист').get(name_pos=request.POST['pos']).id
            spec.email = request.POST['email']
            spec.save()
            context = {
                'curr_user': user,
                'curr_group': group,
                'success': 'Данные успешно изменены'
            }
            return render(request, 'start.html', context)
        else:
            spec = specialists.objects.get(id=id)
            mo = MunObr.objects.all()
            pos = Position.objects.filter(type_pos='Специалист')
            context = {
                'spec': spec,
                'mo': mo,
                'pos': pos,
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'edit_spec.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def add_cert(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        if request.method == 'POST':
            id_max = certified.objects.filter(period_id=period).aggregate(max=Max('id')).get('max')
            code = certified.objects.get(id=id_max).att_code
            code_max = code[8:]
            newcode = int(code_max)+randint(0, 50)
            today = str(date.today())
            att_code = today[:8]+str(newcode)
            new_cert = certified()
            new_cert.period_id=period
            new_cert.att_code = att_code
            new_cert.MO_id = MunObr.objects.get(name_MO=request.POST['mo']).id
            new_cert.FIO = request.POST['fio'].strip()
            new_cert.Organization = request.POST['org']
            new_cert.Position_id = Position.objects.filter(type_pos='Аттестуемый').get(name_pos=request.POST['pos']).id
            new_cert.Category_id = AttCategories.objects.get(type_cat=request.POST['category']).id
            new_cert.att_form_id = request.POST['attform']
            new_cert.save()
            mo = MunObr.objects.all()
            pos = Position.objects.filter(type_pos="Аттестуемый")
            cat = AttCategories.objects.all()
            context = {
                'curr_user': user,
                'curr_group': group,
                'title': 'Аттестуемый '+att_code+' успешно добавлен',
                'period': per,
                'mo': mo,
                'pos': pos,
                'cat': cat
            }
            return render(request, 'import/add_cert.html', context)
        else:
            attforms = AttForm.objects.all()
            mo = MunObr.objects.all()
            pos = Position.objects.filter(type_pos='Аттестуемый')
            cat = AttCategories.objects.all()
            context = {
                'curr_user': user,
                'curr_group': group,
                'period': per,
                'attforms': attforms,
                'mo': mo,
                'pos': pos,
                'cat': cat
            }
            return render(request, 'import/add_cert.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def add_delegate(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        if request.method == 'POST':
            new_spec = delegates()
            new_spec.period_id = period
            new_spec.MO_id = MunObr.objects.get(name_MO=request.POST['mo']).id
            new_spec.FIO = request.POST['fio']
            new_spec.Organization = request.POST['org']
            new_spec.email = request.POST['email']
            new_spec.save()
            mo = MunObr.objects.all()
            context = {
                'curr_user': user,
                'curr_group': group,
                'period': per,
                'mo': mo,
                'title': 'Уполномоченный '+request.POST['fio']+' успешно добавлен'
            }
            return render(request, 'import/add_delegate.html', context)
        else:
            mo = MunObr.objects.all()
            context = {
                'curr_user': user,
                'curr_group': group,
                'period': per,
                'mo': mo,
            }
            return render(request, 'import/add_delegate.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def add_spec(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        if request.method == 'POST':
            new_spec = specialists()
            new_spec.period_id = period
            new_spec.MO_id = MunObr.objects.get(name_MO=request.POST['mo']).id
            new_spec.FIO = request.POST['fio']
            new_spec.Organization = request.POST['org']
            new_spec.Position_id = Position.objects.filter(type_pos='Специалист').get(name_pos=request.POST['pos']).id
            new_spec.email = request.POST['email']
            new_spec.save()
            mo = MunObr.objects.all()
            pos = Position.objects.filter(type_pos='Специалист')
            context = {
                'curr_user': user,
                'curr_group': group,
                'period': per,
                'mo': mo,
                'pos': pos,
                'title': 'Специалист '+request.POST['fio']+' успешно добавлен'
            }
            return render(request, 'import/add_spec.html', context)
        else:
            mo = MunObr.objects.all()
            pos = Position.objects.filter(type_pos='Специалист')
            context = {
                'curr_user': user,
                'curr_group': group,
                'period': per,
                'mo': mo,
                'pos': pos,
            }
            return render(request, 'import/add_spec.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_cert(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        fio = certified.objects.get(id=id).FIO
        certified.objects.get(id=id).delete()
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Аттестуемый '+fio+' был успешно удален'
        }
        return render(request, 'import/success.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_target(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        rec = targets_mo.objects.get(id=id)
        id_del = rec.delegate_id
        rec.delete()
        context = {
            'curr_user': user,
            'curr_group': group,
        }
        return target_mo(request, id_del)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_delegate(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        fio = delegates.objects.get(id=id).FIO
        delegates.objects.get(id=id).delete()
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Уполномоченный '+fio+' был успешно удален'
        }
        return render(request, 'import/success.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_spec(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        fio = specialists.objects.get(id=id).FIO
        specialists.objects.get(id=id).delete()
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Специалист '+fio+' был успешно удален'
        }
        return render(request, 'import/success.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def check_list(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            period = request.POST['period']
            typ = request.POST['type']
            if typ == 'Аттестуемые':
                return cert_post_list(request, period)
            elif typ == 'Специалисты':
                return spec_post_list(request, period)
            else:
                return delegate_post_list(request, period)
        else:
            pers = periods.objects.all().order_by('-id')
            count = pers.count()
            context = {
                'curr_user': user,
                'curr_group': group,
                'count': count,
                'periods': pers
            }
            return render(request, 'import/choose_period.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def find_results(request, period, type):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        fio = request.POST['FIO']
        type_iri = uri_to_iri(type)
        results = Summary_table.objects.filter(period_id=period).filter(type_MO=type_iri).filter(FIO_expert__contains=fio)
        count = results.count()
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': type_iri,
            'results': results,
            'period': per,
            'count': count
        }
        return render(request, 'tables/all_results.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def find_cert(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        fio = request.POST['FIO']
        cert = certified.objects.filter(period_id=period).filter(FIO__contains=fio)
        count = cert.count()
        context = {
            'curr_user': user,
            'curr_group': group,
            'period':per,
            'cert': cert,
            'count': count
        }
        return render(request, 'import/cert_list.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def find_criterias(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        value = request.POST['value']
        field = request.POST['field']
        period = request.POST['per']
        crit_exp = criteria_export.objects.filter(period_id=period)
        if field == 'МО специалиста':
            find = MunObr.objects.filter(name_MO__contains=value).distinct()
            list_mo = []
            for el in find:
                list_mo.append(el)
            criterias = crit_exp.filter(MO_spec_id__in=MunObr.objects.filter(name_MO__in=list_mo))
        elif field == 'ФИО специалиста':
            criterias = crit_exp.filter(FIO_spec__contains=value)
        elif field == 'МО аттестуемого':
            find = MunObr.objects.filter(name_MO__contains=value).distinct()
            list_mo = []
            for el in find:
                list_mo.append(el)
            criterias = crit_exp.filter(MO_att_id__in=MunObr.objects.filter(name_MO__in=list_mo))
        else:
            criterias = crit_exp.filter(FIO_att__contains=value)
        try:
            count = criterias.count()
        except BaseException:
            count = 0
        context = {
            'find': 'yes',
            'curr_user': user,
            'curr_group': group,
            'criterias': criterias,
            'count': count
        }
        return render(request, 'tables/all_criterias.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def cert_post_list(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        object_list = certified.objects.filter(period_id=period)
        count = object_list.count()
        try:
            ord_by = request.GET.get('order_by')
            object_list = object_list.order_by(ord_by)
        except:
            pass
        paginator = Paginator(object_list, 50)
        try:
            page = request.GET.get('page')
        except BaseException:
            page = 1
        try:
            posts = paginator.page(page)
        except PageNotAnInteger:
            # Если страница не является целым числом, поставим первую страницу
            posts = paginator.page(1)
        except EmptyPage:
            # Если страница больше максимальной, доставить последнюю страницу результатов
            posts = paginator.page(paginator.num_pages)
        mo = certified.objects.filter(period_id=period).order_by('MO_id').values('MO_id').distinct()
        list_mo = []
        for el in mo:
            name = MunObr.objects.get(id=el['MO_id']).name_MO
            list_mo.append(name)
        pos = certified.objects.filter(period_id=period).order_by('Position_id').values('Position_id').distinct()
        list_pos = []
        for el in pos:
            name = Position.objects.get(id=el['Position_id']).name_pos
            list_pos.append(name)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        context = {
            'curr_user': user,
            'curr_group': group,
            'period': per,
            'cert': posts,
            'mo': list_mo,
            'pos': list_pos,
            'count': count,
            'order_by': ord_by
        }
        return render(request, 'import/cert_list.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_cert_list(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        del_cert = certified.objects.filter(period_id=period)
        for el in del_cert:
            el.delete()
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Список аттестуемых успешно удален'
        }
        return render(request, 'import/success.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def find_delegate(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        per = period.objects.get(id=period)
        field = request.POST['field']
        value = request.POST['value']
        if field == 'ФИО':
            delegs = delegates.objects.filter(period_id=period).filter(FIO__contains=value)
        else:
            delegs = delegates.objects.filter(period_id=period).filter(email__contains=value)
        count = delegs.count()
        context = {
            'curr_user': user,
            'curr_group': group,
            'period': per,
            'delegates': delegs,
            'count': count,
        }
        return render(request, 'import/delegates_list.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def find_spec(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        field = request.POST['field']
        value = request.POST['value']
        if field == 'ФИО':
            specs = specialists.objects.filter(period_id=period).filter(FIO__contains=value)
        else:
            specs = specialists.objects.filter(period_id=period).filter(email__contains=value)
        count = specs.count()
        context = {
            'curr_user': user,
            'curr_group': group,
            'period': per,
            'specs': specs,
            'count': count,
        }
        return render(request, 'import/specs_list.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def target_spec(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        targs = targets.objects.filter(delegate_id=id)
        count = targs.count()
        deleg = delegates.objects.get(id=id)
        if count == 0:
            title = 'Не найдено назначений от уполномоченного: ' + deleg.FIO
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': title,
            }
            return render(request, 'import/success.html', context)
        list_spec = []
        list_cert = []
        for el in targs:
            list_spec.append(el.spec_id)
            list_cert.append(el.cert_id)
        specs = specialists.objects.filter(id__in=list_spec)
        certs = certified.objects.filter(id__in=list_cert)
        context = {
            'curr_user': user,
            'curr_group': group,
            'targs': targs,
            'count': count,
            'specs': specs,
            'certs': certs,
            'deleg': deleg
        }
        return render(request, 'import/target_spec.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def targets_criteria(request, specid, certid):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        crits = cards_criteria.objects.filter(cert_id=certid).filter(spec_id=specid)
        if group == 'Уполномоченный':
            crits = crits.filter(from_operator=False)
        if request.method == 'POST':
            crits.update(to_deleg=True)
            deleg = delegates.objects.get(MO_id=certified.objects.get(id=certid).MO_id)
            title = 'Замечания успешно предоставлены уполномоченному ' + deleg.FIO
            context = {
                'curr_user': user,
                'curr_group': group,
                'title': title,
            }
            return render(request, 'import/success.html', context)
        count = crits.count()
        context = {
            'curr_user': user,
            'curr_group': group,
            'crits': crits,
            'count': count,
            'certid': certid
        }
        return render(request, 'import/target_criteria.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def targets_operator(request, certid):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        cert = certified.objects.get(id=certid)
        crits = cards_criteria.objects.filter(cert_id=certid)
        list_specs = []
        for el in crits:
            list_specs.append(el.spec_id)
        specs = specialists.objects.filter(id__in=list_specs)
        if group == 'Уполномоченный':
            crits = crits.filter(from_operator=False)
        if request.method == 'POST':
            crits.update(to_deleg=True)
            res = request.POST['res']
            rec = certified.objects.get(id=certid)
            rec.result = res
            rec.save()
            try:
                deleg = delegates.objects.get(MO_id=certified.objects.get(id=certid).MO_id)
                title = 'Замечания успешно предоставлены уполномоченному ' + deleg.FIO
            except BaseException:
                title = 'Готово'
            per_id = certified.objects.get(id=certid).period_id
            return cert_post_list(request, per_id)
        count = crits.count()
        list_result=['Нет результата', 'Установить', 'Отказать']
        context = {
            'curr_user': user,
            'curr_group': group,
            'crits': crits,
            'specs': specs,
            'count': count,
            'cert': cert,
            'res': list_result,
        }
        return render(request, 'import/target_criteria.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def new_crit_operator(request, certid):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            spec = specialists.objects.filter(email=request.user.email).latest('id')
        except BaseException:
            context = {
                'curr_user': user,
                'curr_group': group,
                'title': 'Специалист не найден'
            }
            return render(request, 'import/success.html', context)
        crits = cards_criteria.objects.filter(cert_id=certid)
        cert = certified.objects.get(id=certid)
        criterias = criteria.objects.all()
        if request.method == 'POST':
            inter = crits.latest('id').inter_id
            new_rec = cards_criteria()
            new_rec.inter_id = inter
            new_rec.spec_id = spec.id
            new_rec.cert_id = certid
            new_rec.criteria_id = request.POST['crit']
            new_rec.info = request.POST['info']
            new_rec.to_deleg = False
            new_rec.from_operator = True
            new_rec.save()
            id_card = new_rec.id
            period = specialists.objects.get(id=spec.id).period_id
            new_rec = criteria_export()
            new_rec.period_id = period
            new_rec.id_cards_criteria_id = id_card
            new_rec.MO_spec_id = spec.MO_id
            new_rec.FIO_spec = spec.FIO
            new_rec.Position_spec_id = spec.Position_id
            new_rec.MO_att_id = cert.MO_id
            new_rec.FIO_att = cert.FIO
            new_rec.Position_att_id = cert.Position_id
            new_rec.Category_id = cert.Category_id
            new_rec.Result = intermediate.objects.get(id=inter).result
            new_rec.criteria_id = request.POST['crit']
            new_rec.info = request.POST['info']
            new_rec.save()
            request.method = "GET"
            return targets_operator(request, certid)
        context = {
            'curr_user': user,
            'curr_group': group,
            'criterias': criterias,
            'spec': spec,
            'cert': cert
        }
        return render(request, 'import/add_crit.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def change_targets_criteria(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        rec_id = request.POST['rec_id']
        new_info = request.POST['new_info']
        specid = request.POST['specid']
        certid = request.POST['certid']
        cards_criteria.objects.filter(id=rec_id).update(info=new_info)
        criteria_export.objects.filter(id_cards_criteria=rec_id).update(info=new_info)
        request.method = "GET"
        return targets_criteria(request, specid, certid)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')

def target_mo(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            mos = request.POST.getlist('mo')
            for el in mos:
                new_rec = targets_mo()
                new_rec.MO_id = MunObr.objects.get(name_MO=el).id
                new_rec.delegate_id=id
                new_rec.save()
            return HttpResponse('<meta http-equiv="refresh" content="0; URL=/target_mo/'+str(id)+'">')
        else:
            mo = MunObr.objects.all()
            targs = targets_mo.objects.filter(delegate_id=id)
            count = targs.count()
            deleg = delegates.objects.get(id=id)
            if count == 0:
                title = 'Не найдено назначений для уполномоченного: ' + deleg.FIO
                context = {
                'curr_user': user,
                'curr_group': group,
                'title': title,
                'id': id,
                'mo': mo   
                }
                return render(request, 'import/create_target_mo.html', context)
            context = {
                'curr_user': user,
                'curr_group': group,
                'targs': targs,
                'count': count,
                'deleg': deleg
            }
            return render(request, 'import/target_mo.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def change_target_mo(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            mos = request.POST.getlist('mo')
            for el in mos:
                new_rec = targets_mo()
                new_rec.MO_id = MunObr.objects.get(name_MO=el).id
                new_rec.delegate_id=id
                new_rec.save()
            return HttpResponse('<meta http-equiv="refresh" content="0; URL=/target_mo/'+str(id)+'">')
        else:
            find = targets_mo.objects.filter(delegate_id=id)
            list_mo = []
            for el in find:
                list_mo.append(el.MO_id)
            mo = MunObr.objects.exclude(id__in=list_mo)
            deleg = delegates.objects.get(id=id)
            title = 'Добавить МО для уполномоченного:'+deleg.FIO
            context = {
                'curr_user': user,
                'curr_group': group,
                'mo': mo,
                'id': id,
                'title': title
            }
            return render(request, 'import/create_target_mo.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delegate_post_list(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        object_list = delegates.objects.filter(period_id=period)
        try:
            ord_by = request.GET.get('order_by')
            object_list = object_list.order_by(ord_by)
        except BaseException:
            pass
        count = object_list.count()
        paginator = Paginator(object_list, 50)
        try:
            page = request.GET.get('page')
        except BaseException:
            page = 1
        try:
            posts = paginator.page(page)
        except PageNotAnInteger:
            # Если страница не является целым числом, поставим первую страницу
            posts = paginator.page(1)
        except EmptyPage:
            # Если страница больше максимальной, доставить последнюю страницу результатов
            posts = paginator.page(paginator.num_pages)
        mo = delegates.objects.filter(period_id=period).order_by('MO_id').values('MO_id').distinct()
        list_mo = []
        for el in mo:
            name = MunObr.objects.get(id=el['MO_id']).name_MO
            list_mo.append(name)
        context = {
            'curr_user': user,
            'curr_group': group,
            'period': per,
            'delegates': posts,
            'mo': list_mo,
            'count': count,
            'order_by': ord_by,
        }
        return render(request, 'import/delegates_list.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def spec_post_list(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        object_list = specialists.objects.filter(period_id=period)
        try:
            ord_by = request.GET.get('order_by')
            object_list = object_list.order_by(ord_by)
        except BaseException:
            pass
        count = object_list.count()
        paginator = Paginator(object_list, 50)
        try:
            page = request.GET.get('page')
        except BaseException:
            page = 1
        try:
            posts = paginator.page(page)
        except PageNotAnInteger:
            # Если страница не является целым числом, поставим первую страницу
            posts = paginator.page(1)
        except EmptyPage:
            # Если страница больше максимальной, доставить последнюю страницу результатов
            posts = paginator.page(paginator.num_pages)
        mo = specialists.objects.filter(period_id=period).order_by('MO_id').values('MO_id').distinct()
        list_mo = []
        for el in mo:
            name = MunObr.objects.get(id=el['MO_id']).name_MO
            list_mo.append(name)
        pos = specialists.objects.filter(period_id=period).order_by('Position_id').values('Position_id').distinct()
        list_pos = []
        for el in pos:
            name = Position.objects.get(id=el['Position_id']).name_pos
            list_pos.append(name) 
        context = {
            'curr_user': user,
            'curr_group': group,
            'period': per,
            'specs': posts,
            'mo': list_mo,
            'pos': list_pos,
            'count': count,
            'order_by': ord_by,
        }
        return render(request, 'import/specs_list.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_spec_list(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        del_spec = specialists.objects.filter(period_id=period)
        for el in del_spec:
            el.delete()
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Список специалистов успешно удален'
        }
        return render(request, 'import/success.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_delegate_list(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        del_spec = delegates.objects.filter(period_id=period)
        for el in del_spec:
            el.delete()
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Список уполномоченных успешно удален'
        }
        return render(request, 'import/success.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_all(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        intermediate.objects.all().delete()
        cards_criteria.objects.all().delete()
        Summary_table.objects.all().delete()
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Таблицы успешно очищены'
        }
        return render(request, 'import/success.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def new_criteria(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if request.method == 'POST':
            spec = intermediate.objects.get(id=id).spec_id
            id_period = periods.objects.get(id=specialists.objects.get(id=spec).period_id).id
            cert = intermediate.objects.get(id=id).cert_id
            crit = criteria.objects.get(name_criteria=request.POST['criteria']).id
            att_code = certified.objects.get(id=cert).att_code
            new_rec = cards_criteria()
            new_rec.inter_id = id
            new_rec.spec_id = spec
            new_rec.cert_id = cert
            new_rec.criteria_id = crit
            new_rec.info = request.POST['info']
            new_rec.save()
            rec_id = new_rec.id
            new_crit = criteria_export()
            new_crit.period_id=id_period
            new_crit.id_cards_criteria_id = rec_id
            new_crit.MO_spec_id = specialists.objects.get(id=spec).MO_id
            new_crit.FIO_spec = specialists.objects.get(id=spec).FIO
            new_crit.Position_spec_id = specialists.objects.get(id=spec).Position_id
            new_crit.MO_att_id = certified.objects.get(id=cert).MO_id
            new_crit.FIO_att = certified.objects.get(id=cert).FIO
            new_crit.Position_att_id = certified.objects.get(id=cert).Position_id
            new_crit.Category_id = certified.objects.get(id=cert).Category_id
            new_crit.Result = intermediate.objects.get(id=id).result
            new_crit.criteria_id = crit
            new_crit.info = request.POST['info']
            new_crit.save() 
            info = 'Замечание по критерию '+ request.POST['criteria'] + ' для карты ' \
                    'аттестуемого '+att_code+' успешно добавлено'
            lvls = ExpLevel.objects.all()
            mail = request.user.email
            spec_id = specialists.objects.filter(email=mail).latest('id').id
            count = intermediate.objects.filter(spec_id=spec_id).count()
            cards = intermediate.objects.filter(spec_id=spec_id).order_by("-date_add")
            lvls = ExpLevel.objects.all()
            voc_atts = {}
            for el in cards:
                id = el.cert_id
                voc_atts[id] = certified.objects.filter(id=id).values('att_code').first()
            context = {
                'curr_user': user,
                'curr_group': group,
                'count': count,
                'cards': cards,
                'voc': voc_atts,
                'lvls': lvls,
                'info': info
            }
            return render(request, 'cards_spec.html', context)
        else:
            crit = criteria.objects.all()
            id_cert = intermediate.objects.get(id=id).cert_id
            att_code = certified.objects.get(id=id_cert).att_code
            context = {
                'curr_user': user,
                'curr_group': group,
                'att_code': att_code,
                'criteria': crit
            }
            return render(request, 'criteria/new_criteria.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def show_crit_recs(request, id_cert):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        mail = request.user.email
        spec_id = specialists.objects.filter(email=mail).latest('id').id
        count = cards_criteria.objects.filter(spec_id=spec_id).filter(cert_id=id_cert).count()
        if count == 0:
            context = {
                'curr_user': user,
                'curr_group': group,
                'count': count
            }
            return render(request, 'criteria/show_crit_recs.html', context)
        else:
            records = cards_criteria.objects.filter(spec_id=spec_id).filter(cert_id=id_cert)
            att_code = certified.objects.get(id=id_cert).att_code
            title = 'Замечания по критериям для аттестуемого '+att_code
            context = {
                'curr_user': user,
                'curr_group': group,
                'records': records,
                'count': count,
                'title': title
            }
            return render(request, 'criteria/show_crit_recs.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def show_crit_cards(request, id_inter):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        count = cards_criteria.objects.filter(inter_id=id_inter).count()
        if count == 0:
            context = {
                'curr_user': user,
                'curr_group': group,
                'count': count
            }
            return render(request, 'criteria/show_crit_cards.html', context)
        else:
            records = cards_criteria.objects.filter(inter_id=id_inter)
            title = 'Замечания по критериям'
            context = {
                'curr_user': user,
                'curr_group': group,
                'records': records,
                'count': count,
                'title': title
            }
            return render(request, 'criteria/show_crit_cards.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_crit_rec(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        id_cert = cards_criteria.objects.get(id=id).cert_id
        cards_criteria.objects.get(id=id).delete()
        criteria_export.objects.get(id_cards_criteria=id).delete()
        return show_crit_recs(request, id_cert)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_crit_target(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        id_cert = cards_criteria.objects.get(id=id).cert_id
        cards_criteria.objects.get(id=id).delete()
        return targets_operator(request, id_cert)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def choose_spec_stats(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            typ = request.POST['type']
            period = request.POST['period']
            return statistics_spec(request, period, typ)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        context = {
            'curr_user': user,
            'curr_group': group,
            'period': per
        }
        return render(request, 'import/choose_statistics_spec.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def statistics_spec(request, period, typ):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        specs = specialists.objects.filter(period_id=period)
        if typ == "Статистика по должностям":
            specs = specs.order_by('Position_id')
            positions = specs.values('Position_id').distinct()
            voc_stats = {}
            for el in positions:
                name = Position.objects.get(id=el['Position_id']).name_pos
                count = specs.filter(Position_id=el['Position_id']).count()
                voc_stats[name] = count
            title = "Статистика по должностям"
        elif typ == "Статистика по МО":
            specs = specs.order_by('MO_id')
            mo = specs.values('MO_id').distinct()
            voc_stats = {}
            for el in mo:
                name = MunObr.objects.get(id=el['MO_id']).name_MO
                count = specs.filter(MO_id=el['MO_id']).count()
                voc_stats[name] = count
            title = "Статистика по МО"
        else:
            return main_stat_spec(request, period)
        context = {
                'curr_user': user,
                'title': title,
                'curr_group': group,
                'voc': voc_stats
        }
        return render(request, 'import/stats.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def choose_cert_stats(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1] 
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        if request.method == 'POST':
            typ = request.POST['type']
            return statistics_cert(request, period, typ)
        context = {
            'curr_user': user,
            'curr_group': group,
            'period': per
        }
        return render(request, 'import/choose_statistics_cert.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def statistics_cert(request, period, typ):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        certs = certified.objects.filter(period_id=period)
        if typ == "Статистика по должностям":
            certs = certs.order_by('Position_id')
            positions = certs.values('Position_id').distinct()
            voc_stats = {}
            for el in positions:
                name = Position.objects.get(id=el['Position_id']).name_pos
                count = certs.filter(Position_id=el['Position_id']).count()
                voc_stats[name] = count
            title = "Статистика по должностям"
        elif typ == "Статистика по МО":
            certs = certs.order_by('MO_id')
            mo = certs.values('MO_id').distinct()
            voc_stats = {}
            for el in mo:
                name = MunObr.objects.get(id=el['MO_id']).name_MO
                count = certs.filter(MO_id=el['MO_id']).count()
                voc_stats[name] = count
            title = "Статистика по МО"
        elif typ == "Общая статистика МО/Должность":
            return main_stat_cert(request, period)
        else:
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            )
            response['Content-Disposition'] = 'attachment; filename=queries_stats.xlsx'
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = 'Данные'
            cell = worksheet.cell(row=1,column=1)
            cell.value = 'Статистические данные'	
            cell = worksheet.cell(row=1,column=2)
            cell.value = 'по заявкам за период "'+per.name_period+'"'
            cell = worksheet.cell(row=2,column=1)
            cell.value = 'Дата начала периода'
            cell = worksheet.cell(row=2,column=2)
            cell.value = per.date_start
            cell = worksheet.cell(row=3,column=1)
            cell.value = 'Дата окончания периода'
            cell = worksheet.cell(row=3,column=2)
            cell.value = per.date_end
            titles = [
                '№',
                'Наименование (МО, ведомства, ОО)',
                'Всего заявок',
                'в том числе на 1КК',
                'в том числе на ВКК',
                'Аттест. Сессия',
                'Эксперт заключ.',
                'Модельн паспорт',
                '1КК',
                'ВКК',
            ]
            row_num = 5
            for col_num, column_title in enumerate(titles, 1):
                cell = worksheet.cell(row=row_num, column=col_num)
                cell.value = column_title
            count_main = 0
            count_1kk = 0
            count_vkk = 0
            count_attses = 0
            count_exp =0
            count_model = 0
            mo = MunObr.objects.filter(type_MO="Муниципалитет").values('name_MO')
            cert = certified.objects.filter(period_id=period)
            for el in mo:
                row_num += 1
                mo_id = MunObr.objects.get(name_MO=el['name_MO']).id
                row = []
                row.append(row_num-5)
                row.append(el['name_MO'])
                try:
                    count = cert.filter(MO_id=mo_id).count()
                except BaseException:
                    count = 0
                count_main += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(Category_id=AttCategories.objects.get(type_cat='первая').id).count()
                except BaseException:
                    count = 0
                count_1kk += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(Category_id=AttCategories.objects.get(type_cat='высшая').id).count()
                except BaseException:
                    count = 0
                count_vkk += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(att_form_id=AttForm.objects.get(name='Аттестационная сессия').id).count()
                except BaseException:
                    count = 0
                count_attses += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(att_form_id=AttForm.objects.get(name='Экспертное заключение').id).count()
                except BaseException:
                    count = 0
                count_exp += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(att_form_id=AttForm.objects.get(name='Модельный паспорт').id).count()
                except BaseException:
                    count = 0
                count_model += count
                row.append(count)
                row.append('0')
                row.append('0')
                for col_num, cell_value in enumerate(row, 1):
                    cell = worksheet.cell(row=row_num, column=col_num)
                    cell.value = cell_value
            row_num += 1
            row = []
            row.append('')
            row.append('Итого по МО')
            row.append(count_main)
            row.append(count_1kk)
            row.append(count_vkk)
            row.append(count_attses)
            row.append(count_exp)
            row.append(count_model)
            row.append('0')
            row.append('0')
            for col_num, cell_value in enumerate(row, 1):
                cell = worksheet.cell(row=row_num, column=col_num)
                cell.value = cell_value
            row_num += 1
            count_gos_main = 0
            count_gos_1kk = 0
            count_gos_vkk = 0
            count_gos_attses = 0
            count_gos_exp =0
            count_gos_model = 0
            mo = MunObr.objects.filter(type_MO="Гос").values('name_MO')
            cert = certified.objects.filter(period_id=period)
            for el in mo:
                row_num += 1
                mo_id = MunObr.objects.get(name_MO=el['name_MO']).id
                row = []
                row.append(row_num-7)
                row.append(el['name_MO'])
                try:
                    count = cert.filter(MO_id=mo_id).count()
                except BaseException:
                    count = 0
                count_gos_main += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(Category_id=AttCategories.objects.get(type_cat='первая').id).count()
                except BaseException:
                    count = 0
                count_gos_1kk += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(Category_id=AttCategories.objects.get(type_cat='высшая').id).count()
                except BaseException:
                    count = 0
                count_gos_vkk += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(att_form_id=AttForm.objects.get(name='Аттестационная сессия').id).count()
                except BaseException:
                    count = 0
                count_gos_attses += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(att_form_id=AttForm.objects.get(name='Экспертное заключение').id).count()
                except BaseException:
                    count = 0
                count_gos_exp += count
                row.append(count)
                try:
                    count = cert.filter(MO_id=mo_id).filter(att_form_id=AttForm.objects.get(name='Модельный паспорт').id).count()
                except BaseException:
                    count = 0
                count_gos_model += count
                row.append(count)
                row.append('0')
                row.append('0')
                for col_num, cell_value in enumerate(row, 1):
                    cell = worksheet.cell(row=row_num, column=col_num)
                    cell.value = cell_value
            row_num += 1
            row = []
            row.append('')
            row.append('Итого по другим министерствам')
            row.append(count_gos_main)
            row.append(count_gos_1kk)
            row.append(count_gos_vkk)
            row.append(count_gos_attses)
            row.append(count_gos_exp)
            row.append(count_gos_model)
            row.append('0')
            row.append('0')
            for col_num, cell_value in enumerate(row, 1):
                cell = worksheet.cell(row=row_num, column=col_num)
                cell.value = cell_value
            row_num += 2
            row = []
            row.append('')
            row.append('Итого')
            row.append(count_main+count_gos_main)
            row.append(count_1kk+count_gos_1kk)
            row.append(count_vkk+count_gos_vkk)
            row.append(count_attses+count_gos_attses)
            row.append(count_exp+count_gos_exp)
            row.append(count_model+count_gos_model)
            row.append('0')
            row.append('0')
            for col_num, cell_value in enumerate(row, 1):
                cell = worksheet.cell(row=row_num, column=col_num)
                cell.value = cell_value
            workbook.save(response)
            return response
        context = {
            'curr_user': user,
            'title': title,
            'curr_group': group,
            'voc': voc_stats
        }
        return render(request, 'import/stats.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def main_stat_cert(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        certs = certified.objects.filter(period_id=period)
        ln = 0
        try:
            mo = request.POST.getlist('mo')
            ln = len(mo)
        except BaseException:
            pass
        if request.method == 'POST' and ln != 0:
            mo = request.POST.getlist('mo')
            pos = request.POST.getlist('pos')
            certs = certs.order_by('MO_id')
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            )
            response['Content-Disposition'] = 'attachment; filename=certified_stats.xlsx'
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = 'Статистика'
            row_num = 1
            cell = worksheet.cell(row=1,column=1)
            cell.value = 'МО'
            for col_num, column_title in enumerate(pos, 1):
                cell = worksheet.cell(row=row_num, column=col_num+1)
                cell.value = column_title
            for el in mo:
                row_num += 1
                mo_id = MunObr.objects.get(name_MO=el).id
                row = []
                row.append(el)
                for el_p in pos:
                    pos_id = Position.objects.filter(type_pos='Аттестуемый').get(name_pos=el_p).id
                    try:
                        count = certs.filter(Position_id=pos_id).filter(MO_id=mo_id).count()
                    except BaseException:
                        count = 0
                    row.append(count)
                for col_num, cell_value in enumerate(row, 1):
                    cell = worksheet.cell(row=row_num, column=col_num)
                    cell.value = cell_value
            workbook.save(response)
            return response
        mo = certs.values('MO_id').distinct().order_by('MO_id')
        list_mo = []
        for el in mo:
            name = MunObr.objects.get(id=el['MO_id']).name_MO
            list_mo.append(name)
        pos = certs.values('Position_id').distinct().order_by('Position_id')
        list_pos = []
        for el in pos:
            name = Position.objects.filter(type_pos='Аттестуемый').get(id=el['Position_id']).name_pos
            list_pos.append(name)
        context = {
            'curr_user': user,
            'curr_group': group,
            'mo': list_mo,
            'pos': list_pos,
            'period': per
        }
        return render(request, 'import/choose_fields_stat.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def main_stat_spec(request, period):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=period)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        specs = specialists.objects.filter(period_id=period)
        ln = 0
        try:
            mo = request.POST.getlist('mo')
            ln = len(mo)
        except BaseException:
            pass
        if request.method == 'POST' and ln != 0:
            mo = request.POST.getlist('mo')
            pos = request.POST.getlist('pos')
            specs = specs.order_by('MO_id')
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            )
            response['Content-Disposition'] = 'attachment; filename=specialists_stats.xlsx'
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = 'Статистика'
            row_num = 1
            cell = worksheet.cell(row=1,column=1)
            cell.value = 'МО'
            for col_num, column_title in enumerate(pos, 1):
                cell = worksheet.cell(row=row_num, column=col_num+1)
                cell.value = column_title
            for el in mo:
                row_num += 1
                mo_id = MunObr.objects.get(name_MO=el).id
                row = []
                row.append(el)
                for el_p in pos:
                    pos_id = Position.objects.filter(type_pos='Специалист').get(name_pos=el_p).id
                    try:
                        count = specs.filter(Position_id=pos_id).filter(MO_id=mo_id).count()
                    except BaseException:
                        count = 0
                    row.append(count)
                for col_num, cell_value in enumerate(row, 1):
                    cell = worksheet.cell(row=row_num, column=col_num)
                    cell.value = cell_value
            workbook.save(response)
            return response
        mo = specs.values('MO_id').distinct().order_by('MO_id')
        list_mo = []
        for el in mo:
            name = MunObr.objects.get(id=el['MO_id']).name_MO
            list_mo.append(name)
        pos = specs.values('Position_id').distinct().order_by('Position_id')
        list_pos = []
        for el in pos:
            name = Position.objects.filter(type_pos='Специалист').get(id=el['Position_id']).name_pos
            list_pos.append(name)
        context = {
            'curr_user': user,
            'curr_group': group,
            'mo': list_mo,
            'pos': list_pos,
            'period': per,
            'spec': 'yes'
        }
        return render(request, 'import/choose_fields_stat.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')



def show_periods(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        count = periods.objects.all().count()
        if count == 0:
            return new_period(request)
        else:
            pers = periods.objects.all().order_by('-id')
            title = 'Периоды'
            context = {
                'curr_user': user,
                'curr_group': group,
                'periods': pers,
                'count': count,
                'title': title
            }
        return render(request, 'periods/show_periods.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def new_period(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        if request.method == 'POST':
            name = request.POST['name']
            if len(name) > 30:
                context = {
                            'curr_user': user,
                            'curr_group': group,
                            'error': 'Слишком длинное название периода (максимум - 30 символов)'
                }
                return render(request, 'periods/new_period.html', context)
            date_start = request.POST['date_start']
            date_end = request.POST['date_end']
            pers = periods.objects.all()
            if pers.count() > 0:
                for el in pers:
                    if el.name_period == name:
                        context = {
                            'curr_user': user,
                            'curr_group': group,
                            'error': 'Период с таким именем уже существует'
                        }
                        return render(request, 'periods/new_period.html', context)
            new = periods()
            new.name_period = name
            new.date_start = date_start
            new.date_end = date_end
            new.save()
            return show_periods(request)
        else:
            context = {
                'curr_user': user,
                'curr_group': group,
            }
            return render(request, 'periods/new_period.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def edit_period(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        try:
            per = periods.objects.get(id=id)
        except BaseException:
            context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период не найден или найдено более двух периодов'
            }
            return render(request, 'import/success.html', context)
        if request.method == 'POST':
            name = request.POST['name']
            date_start = request.POST['date_start']
            date_end = request.POST['date_end']
            per.name = name
            per.date_start = date_start
            per.date_end = date_end
            per.save()
            pers = periods.objects.all().order_by('-id')
            count = pers.count()
            title = 'Период "'+name+'" успешно изменен'
            context = {
                'curr_user': user,
                'curr_group': group,
                'periods': pers,
                'count': count,
                'title': title
            }
            return render(request, 'periods/show_periods.html', context)
        else:
            
            date_start = str(per.date_start)
            date_end = str(per.date_end)
            context = {
                'curr_user': user,
                'curr_group': group,
                'period': per,
                'date_start': date_start,
                'date_end': date_end
            }
            return render(request, 'periods/edit_period.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def delete_period(request, id):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if check_access(request, group):
            context = {
                'curr_user': user,
                'curr_group': group
            }
            return render(request, 'access_denied.html', context)
        name = periods.objects.get(id=id).name_period
        periods.objects.get(id=id).delete()
        pers = periods.objects.all().order_by('-id')
        context = {
            'curr_user': user,
            'curr_group': group,
            'title': 'Период "'+name+'" успешно удален',
            'periods': pers
        }
        return render(request, 'periods/show_periods.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')


def choose_certs_crits(request):
    if request.user.is_authenticated:
        data = get_user_and_group(request)
        user = data[0]
        group = data[1]
        if request.method == "POST":
            list_certs = request.POST.getlist('certs')
            crits = cards_criteria.objects.filter(cert_id__in=list_certs).\
                filter(to_deleg=True).values('cert_id', 'criteria_id', 'info')
            count = crits.count()
            voc_certs = {}
            for el in list_certs:
                c = certified.objects.get(id=el)
                voc_certs[int(el)] = c.att_code+' ('+c.result+')'
            criterias = criteria.objects.all()
            voc_crits = {}
            for el in criterias:
                voc_crits[el.id] = el.name_criteria
            context = {
                'curr_user': user,
                'curr_group': group,
                'voc_certs': voc_certs,
                'voc_crits': voc_crits,
                'crits': crits,
                'count': count
            }
            return render(request, 'import/crit_targets_list.html', context)
        per = periods.objects.latest('id')
        try:
            mo = delegates.objects.filter(period=per.id).get(email=request.user.email).MO_id
        except BaseException:
            return HttpResponse('Не найдены замечания за текущий период')
        certs = certified.objects.filter(period=per.id).filter(MO_id=mo)
        voc_certs = {}
        for el in certs:
            voc_certs[el.id] = el.att_code+' ('+el.result+')'
        count = certs.count()
        context = {
            'curr_user': user,
            'curr_group': group,
            'certs': voc_certs,
            'count': count,
        }
        return render(request, 'import/choose_certs_crits.html', context)
    else:
        return HttpResponse('<meta http-equiv="refresh" content="0; URL=/accounts/login/">')

# Create your views here.

